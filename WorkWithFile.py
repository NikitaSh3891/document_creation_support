import docx
from WorkWithFileCreate.EditFileStyle import getFileStyleCollectionFromPattern, FileFillingText, replaceSymbol
from WorkWithDB import loggingInfo

"""
Класс предназначен для создания файлов с расширением .docx
"""


def editParamUser(param):
    global paramUser
    paramUser = param


def createDocument(fileName, checkedList, idPattern):
    try:
        doc = docx.Document()
        getFileStyleCollectionFromPattern(doc, idPattern)
        for info in checkedList:
            FileFillingText(doc, replaceSymbol(info), idPattern)
        fileExtensionInput = ".docx"
        doc.save('OutputFiles/' + fileName + fileExtensionInput)
        loggingInfo(2, paramUser[0], "Код экспериментов: " + str(checkedList), 2, 1)
        return True
    except:
        return False


def createDocumentWithTitlePage(fileName, nameTitle, checkedList, idPattern):
    try:
        doc = docx.Document('OutputFiles\\' + nameTitle)
        getFileStyleCollectionFromPattern(doc, idPattern)
        for info in checkedList:
            FileFillingText(doc, replaceSymbol(info), idPattern)
        fileExtensionInput = ".docx"
        doc.save('OutputFiles/' + fileName + fileExtensionInput)
        loggingInfo(2, paramUser[0], f"Код титульного листа: {str(nameTitle)}; Код экспериментов: {str(checkedList)}",
                    2, 1)
        return True
    except:
        return False

