import docx
import zipfile
import pathlib

"""
НЕ ИСПОЛЬЗУЮЩИЙСЯ КЛАСС-НАРАБОТКА
Класс предназначен для загрузки содержимого файла в базу данных, в том числе картинок
"""


# https://vc.ru/dev/185015-analiz-dokumentov-word-s-ispolzovaniem-python#:~:text=%D0%A0%D0%B0%D0%B1%D0%BE%D1%82%D0%B0%20%D1%81%20%D0%B8%D0%BB%D0%BB%D1%8E%D1%81%D1%82%D1%80%D0%B0%D1%86%D0%B8%D1%8F%D0%BC%D0%B8

def openFile(fileName):
    doc = docx.Document(fileName)
    print(len(doc.paragraphs))
    for par in doc.paragraphs:
        print(par.style.name, par.text)


def scanImage():
    fileName = '1.docx'
    outputDir = pathlib.Path(f'pic_{fileName}')
    if not outputDir.is_dir():
        outputDir.mkdir()
    with zipfile.ZipFile(fileName) as zf:
        for name in zf.infolist():
            if name.filename.startswith('word/media/'):
                print(name)
                zf.extract(name, outputDir)


# scanImage()
# openFile("../OutputFiles/1.docx")

