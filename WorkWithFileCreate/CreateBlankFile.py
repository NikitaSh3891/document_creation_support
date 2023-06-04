import docx
from docx.shared import Pt, Cm
from datetime import datetime

from WorkWithDB import findNameAndTextDocument
from Converter import convertStringToFile, convertFile, deleteFile
from WorkWithDB import findPositionCellInBlank

"""
Класс предназначен для редактирования бланков
"""


def createEmptyBlank(fileName, statArr):
    documentArr = findNameAndTextDocument(23)
    documentName = documentArr[0][0] + f"_{fileName}"
    convertStringToFile(documentArr[0][1].replace('"', '\''), documentName)
    convertFile(documentName, ".rtf", ".docx")
    deleteFile(documentName, ".rtf")
    fillBlank(documentName, statArr, fileName)


def fillBlank(fileNameBlank, statArr, fileName):
    blankFile = open("OutputFiles\\" + fileNameBlank + ".docx", 'rb')
    doc = docx.Document(blankFile)
    positionArr = findPositionCellInBlank(2)
    for table in doc.tables:
        count = 0
        for positionStr in positionArr:
            elementPos = str(positionStr[0]).split(" ")
            addValueCell(table, int(elementPos[0]), int(elementPos[1]), str(int(statArr[count])))
            count += 1

    doc.add_paragraph("Имя файла: " + fileNameBlank)
    doc.add_paragraph("Дата проверки: " + str(datetime.now().date()))
    parImage = doc.add_paragraph()
    run = parImage.add_run()
    run.add_picture('OutputFiles\\Результаты проверки файла\\' + fileName + ".png", width=Cm(15))
    try:
        doc.save('OutputFiles\\Результаты проверки файла\\' + fileNameBlank + ".docx")
    except PermissionError:
        from Windows.QtWindows.errWindow import editText
        editText("Ошибка! Пожалуйста, закройте файл")
    blankFile.close()
    deleteFile(fileNameBlank, ".docx")
    deleteFile("\\Результаты проверки файла\\" + fileName, ".png")


def addValueCell(table, column, row, text):
    cell = table.cell(column, row)
    cell.paragraphs[0].add_run(text)
    rc = cell.paragraphs[0].runs[0]
    rc.font.name = 'Times New Roman'
    rc.font.size = Pt(14)
