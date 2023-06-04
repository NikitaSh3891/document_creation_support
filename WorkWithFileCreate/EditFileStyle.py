from docx.shared import Pt, Cm
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT, WD_BREAK
from docx.shared import RGBColor
from docx.enum.style import WD_STYLE_TYPE
from distutils.util import strtobool

from WorkWithDB import findFileStyleCollection, findStyleCollectionFromPattern, findDataExperiment
from Converter import convertStringToImage

"""
Класс предназначен для редактирование стилей файла .docx
"""


def pageMarginTest(doc, nameStyle, idPattern):
    styleArr = findFileStyleCollection(nameStyle, idPattern)
    sections = doc.sections
    for section in sections:
        section.left_margin = Cm(int(replaceSymbol(styleArr[0])))
        section.right_margin = Cm(int(replaceSymbol(styleArr[1])))
        section.top_margin = Cm(int(replaceSymbol(styleArr[2])))
        section.bottom_margin = Cm(int(replaceSymbol(styleArr[3])))


def addText(doc, style, text, idPattern):
    par = doc.add_paragraph(style=style)
    # if style == "Heading 1":
    #     par.add_run().add_break(WD_BREAK.PAGE)
    run = par.add_run(text)
    styleArr = findFileStyleCollection(style, idPattern)
    run.font.name = replaceSymbol(styleArr[0])
    run.font.size = Pt(int(replaceSymbol(styleArr[1])))
    color = replaceSymbol(styleArr[10]).split(" ")
    run.font.color.rgb = RGBColor(int(color[0]), int(color[1]), int(color[2]))


def FileFillingText(doc, idExperiment, idPattern):
    dataExperiment = findDataExperiment(idExperiment)
    for data in dataExperiment:
        nameStyle = replaceSymbol(data[1])
        text = replaceSymbol(data[0])
        if nameStyle == "Image":
            addImage(doc, nameStyle, text)
        elif nameStyle == "Table":
            addTable(doc, nameStyle, text, idPattern)
        else:
            addText(doc, nameStyle, text, idPattern)


def replaceSymbol(text):
    text = str(text)
    text = text.replace('[', '').replace(']', '').replace('(', '').replace(')', '').replace('\'', '').replace(',', '')
    return text


def editFileStyle(doc, nameStyle, idPattern):
    styleArr = findFileStyleCollection(nameStyle, idPattern)
    style = doc.styles
    style[nameStyle].paragraph_format.line_spacing = float(replaceSymbol(styleArr[2]))
    match replaceSymbol(styleArr[3]):
        case "Ширине":
            style[nameStyle].paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY
        case "Центру":
            style[nameStyle].paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
        case "Лево":
            style[nameStyle].paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT
        case "Право":
            style[nameStyle].paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.RIGHT
    style[nameStyle].paragraph_format.first_line_indent = Cm(float(replaceSymbol(styleArr[4])))
    match replaceSymbol(styleArr[5]):
        case "Жирный":
            style[nameStyle].font.bold = True
            style[nameStyle].font.italic = False
            style[nameStyle].font.underline = False
        case "Курсив":
            style[nameStyle].font.bold = False
            style[nameStyle].font.italic = True
            style[nameStyle].font.underline = False
        case "Подчеркивание":
            style[nameStyle].font.italic = False
            style[nameStyle].font.bold = False
            style[nameStyle].font.underline = True
        case "ЖирныйИКурсив":
            style[nameStyle].font.italic = True
            style[nameStyle].font.bold = True
            style[nameStyle].font.underline = False
        case "ЖирныйИПодчеркивание":
            style[nameStyle].font.italic = False
            style[nameStyle].font.bold = True
            style[nameStyle].font.underline = True
        case "КурсивИПодчеркивание":
            style[nameStyle].font.italic = True
            style[nameStyle].font.bold = False
            style[nameStyle].font.underline = True
        case "ЖирныйИКурсивИПодчеркивание":
            style[nameStyle].font.italic = True
            style[nameStyle].font.bold = True
            style[nameStyle].font.underline = True
        case "Обычный":
            style[nameStyle].font.italic = False
            style[nameStyle].font.bold = False
            style[nameStyle].font.underline = False
    style[nameStyle].paragraph_format.space_after = Cm(float(replaceSymbol(styleArr[6])))
    style[nameStyle].paragraph_format.space_before = Cm(float(replaceSymbol(styleArr[7])))
    style[nameStyle].paragraph_format.keep_with_next = strtobool(replaceSymbol(styleArr[8]))
    style[nameStyle].font.all_caps = strtobool(replaceSymbol(styleArr[9]))


def getFileStyleCollectionFromPattern(doc, idPattern):
    global nameStyle
    try:
        stylesArr = findStyleCollectionFromPattern(idPattern)
        for nameStyle in stylesArr:
            nameStyle = replaceSymbol(nameStyle)
            if nameStyle == "Common":
                pageMarginTest(doc, nameStyle, idPattern)
            elif nameStyle == "Table" or nameStyle == "Image":
                addFileStyle(doc, nameStyle)
                editFileStyle(doc, nameStyle, idPattern)
            else:
                editFileStyle(doc, nameStyle, idPattern)
    except:
        doc.styles.add_style(nameStyle, WD_STYLE_TYPE.PARAGRAPH)
        getFileStyleCollectionFromPattern(doc, idPattern)


def addFileStyle(doc, nameStyle):
    doc.styles.add_style(nameStyle, WD_STYLE_TYPE.PARAGRAPH)


def addImage(doc, style, text):
    doc.add_paragraph()
    parImage = doc.add_paragraph(style=style)
    run = parImage.add_run()
    run.add_picture(convertStringToImage(text[1:].encode('utf-8')), width=Cm(15))


def addTable(doc, style, text, idPattern):
    tableText = text.split("|")
    table = doc.add_table(rows=int(tableText[0]), cols=int(tableText[1]), style="Table Grid")
    count = 2
    styleArr = findFileStyleCollection(style, idPattern)
    for row in range(int(tableText[0])):
        for col in range(int(tableText[1])):
            cell = table.cell(row, col)
            par = cell.paragraphs[0]
            par.style = doc.styles[style]
            run = par.add_run(tableText[count])
            run.font.name = replaceSymbol(styleArr[0])
            run.font.size = Pt(int(replaceSymbol(styleArr[1])))
            color = replaceSymbol(styleArr[10]).split(" ")
            run.font.color.rgb = RGBColor(int(color[0]), int(color[1]), int(color[2]))
            count += 1
    table.autofit = True
