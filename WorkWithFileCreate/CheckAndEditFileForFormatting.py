import docx
from docx.shared import Pt, Cm
import os
import re

from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from thefuzz import fuzz
from docx.shared import RGBColor
from distutils.util import strtobool

from Converter import convertFileToDOCX, deleteDOCXFile
from WorkWithDB import findStyleCollectionFromPattern, findFileStyleCollection
from WorkWithFileCreate.EditFileStyle import replaceSymbol

"""
Класс предназначен для создания исправленного файла, прошедшего проверку
"""


reValue = "[a-z]|[A-Z]|[а-я]|[А-Я]|[ё|Ё]"
dictionaryOfQuotationMarks = ["'", "\"", "«", "»"]


def filterElement(document):
    res = []
    for element in document.elements:
        if 'paragraph' in str(element):
            res.append(element)
        elif 'table' in str(element):
            for row in element.rows:
                for cell in row.cells:
                    res.append(cell.paragraphs)
    return res


def findNumberPage(text, paragraphs):
    # список для сбора номеров страниц, где встретилась строка
    pages = []
    # счетчик для номеров страниц
    numberPage = 1
    for paragraph in paragraphs:
        if type(paragraph) == list:
            # проверка на наличие разрывов страниц внутри таблиц
            text_in_table = [p.text for p in paragraph]
            text_in_table = ''.join(text_in_table)
            for p in paragraph:
                for run in p.runs:
                    # проверка на мягкий разрыва страницы
                    if 'lastRenderedPageBreak' in run._element.xml:
                        numberPage += 1
                    # проверка на жесткий разрыв страницы
                    elif 'w:br' in run._element.xml and 'type="page"' in run._element.xml:
                        numberPage += 1
            # к этому моменту известен актуальный номер страницы
            if len(text_in_table) >= len(text) - 10:
                res = fuzz.partial_ratio(text.lower(), text_in_table.lower())
                if res >= 97:
                    # если строчка найдена, добавляем в список номер страницы
                    pages.append(numberPage)
        else:
            # проверка разрывов в абзацах вне таблиц
            for run in paragraph.runs:
                if 'lastRenderedPageBreak' in run._element.xml:
                    numberPage += 1
                elif 'w:br' in run._element.xml and 'type="page"' in run._element.xml:
                    numberPage += 1
            if len(paragraph.text) >= len(text):
                res = fuzz.partial_ratio(text.lower(), paragraph.text.lower())
                if res >= 97:
                    pages.append(numberPage)
    return ', '.join(map(str, pages))


# Число букв, выводимое в файл для определения места ошибки
def editNumCharacter(numCharacter):
    global num
    num = numCharacter


# Исправление отступов документа
def pageMarginTest(docInput, nameStyle, idPattern):
    styleArr = findFileStyleCollection(nameStyle, idPattern)
    sections = docInput.sections
    for section in sections:
        if round(section.left_margin.cm, 2) != float(replaceSymbol(styleArr[0])):
            section.left_margin = Cm(int(replaceSymbol(styleArr[0])))
        if round(section.right_margin.cm, 2) != float(replaceSymbol(styleArr[1])):
            section.right_margin = Cm(int(replaceSymbol(styleArr[1])))
        if round(section.top_margin.cm, 2) != float(replaceSymbol(styleArr[2])):
            section.top_margin = Cm(int(replaceSymbol(styleArr[2])))
        if round(section.bottom_margin.cm, 2) != float(replaceSymbol(styleArr[3])):
            section.bottom_margin = Cm(int(replaceSymbol(styleArr[3])))


# Исправление шрифта
def fontNameTest(parRuns, trueStyle, style, paragraph):
    if parRuns.font.name is None:
        if style[paragraph.style.name].font.name != replaceSymbol(trueStyle[1]):
            if paragraph.style.base_style is not None:
                if style[paragraph.style.base_style.name].font.name != replaceSymbol(trueStyle[1]):
                    parRuns.text = parRuns.text.translate({ord(i): None for i in '\n'})
                    if len(parRuns.text) > 0:
                        style[paragraph.style.name].font.name = replaceSymbol(trueStyle[1])
    else:
        if parRuns.font.name != replaceSymbol(trueStyle[1]):
            parRuns.text = parRuns.text.translate({ord(i): None for i in '\n'})
            if len(parRuns.text) > 0:
                parRuns.font.name = replaceSymbol(trueStyle[1])


# Исправление размера шрифта
def fontSizeTest(parRuns, trueStyle, style, paragraph):
    if parRuns.font.size is None:
        if style[paragraph.style.name].font.size != Pt(int(replaceSymbol(trueStyle[2]))):
            if paragraph.style.base_style is not None:
                if style[paragraph.style.base_style.name].font.size != Pt(int(replaceSymbol(trueStyle[2]))):
                    parRuns.text = parRuns.text.translate({ord(i): None for i in '\n'})
                    if len(parRuns.text) > 0:
                        style[paragraph.style.base_style.name].font.size = Pt(int(replaceSymbol(trueStyle[2])))
    else:
        if parRuns.font.size != Pt(int(replaceSymbol(trueStyle[2]))):
            if parRuns.text != " ":
                parRuns.font.size = Pt(int(replaceSymbol(trueStyle[2])))


# Исправление на полужирный, подчеркнутый и курсивный текст
def fontBoldItalicUnderlineTest(parRuns, trueStyle, style, paragraph):
    trueValueBold = None
    trueValueItalic = None
    trueValueUnderline = None
    match replaceSymbol(trueStyle[6]):
        case "Жирный":
            trueValueBold = True
            trueValueItalic = False
            trueValueUnderline = False
        case "Курсив":
            trueValueBold = False
            trueValueItalic = True
            trueValueUnderline = False
        case "Подчеркивание":
            trueValueBold = False
            trueValueItalic = False
            trueValueUnderline = True
        case "ЖирныйИКурсив":
            trueValueBold = True
            trueValueItalic = True
            trueValueUnderline = False
        case "ЖирныйИПодчеркивание":
            trueValueBold = True
            trueValueItalic = False
            trueValueUnderline = True
        case "КурсивИПодчеркивание":
            trueValueBold = False
            trueValueItalic = True
            trueValueUnderline = True
        case "ЖирныйИКурсивИПодчеркивание":
            trueValueBold = True
            trueValueItalic = True
            trueValueUnderline = True
        case "Обычный":
            trueValueBold = False
            trueValueItalic = False
            trueValueUnderline = False
    if parRuns.font.bold is None:
        if style[paragraph.style.name].font.bold and not trueValueBold:
            parRuns.text = parRuns.text.translate({ord(i): None for i in '\n'})
            if len(parRuns.text) > 0:
                style[paragraph.style.name].font.bold = trueValueUnderline
    else:
        if parRuns.font.bold and not trueValueBold:
            if parRuns.text != " ":
                parRuns.font.bold = trueValueUnderline
    if parRuns.font.italic is None:
        if style[paragraph.style.name].font.italic and not trueValueItalic:
            parRuns.text = parRuns.text.translate({ord(i): None for i in '\n'})
            if len(parRuns.text) > 0:
                style[paragraph.style.name].font.italic = trueValueUnderline
    else:
        if parRuns.font.italic and not trueValueItalic:
            if parRuns.text != " ":
                parRuns.font.italic = trueValueUnderline
    if parRuns.font.underline is None:
        if style[paragraph.style.name].font.underline and not trueValueUnderline:
            parRuns.text = parRuns.text.translate({ord(i): None for i in '\n'})
            if len(parRuns.text) > 0:
                style[paragraph.style.name].font.underline = trueValueUnderline
    else:
        if parRuns.font.underline and not trueValueUnderline:
            if parRuns.text != " ":
                parRuns.font.underline = trueValueUnderline


# Исправление на заглавные буквы
def fontAllCapsTest(parRuns, trueStyle, style, paragraph):
    if parRuns.font.all_caps is None:
        if replaceSymbol(trueStyle[10]) == "False":
            if parRuns.text == parRuns.text.upper():
                if re.search(reValue, parRuns.text) is not None:
                    if len(re.sub('[0-9 ]+', '', parRuns.text)) != 1:
                        numberStartStr = paragraph.text.find(parRuns.text)
                        if paragraph.text[numberStartStr - 1] != " " and \
                                re.search(reValue, paragraph.text[numberStartStr - 1]) is None:
                            if paragraph.text[numberStartStr - 1] not in dictionaryOfQuotationMarks:
                                parRuns.text = parRuns.text.lower()
                        else:
                            isTrue = False
                            numSpace = numberStartStr - 1
                            while numSpace > 0:
                                if paragraph.text[numSpace] in dictionaryOfQuotationMarks:
                                    isTrue = True
                                    break
                                numSpace -= 1
                            if not isTrue:
                                parRuns.text = parRuns.text.lower()
    else:
        if str(parRuns.font.all_caps) != replaceSymbol(trueStyle[10]):
            if re.search(reValue, parRuns.text) is not None:
                parRuns.font.all_caps = strtobool(replaceSymbol(trueStyle[10]))


# Исправление на цвет текста
def fontColorTest(parRuns, trueStyle, style, paragraph):
    if parRuns.font.color.rgb is not None:
        color = replaceSymbol(trueStyle[11]).split(" ")
        if parRuns.font.color.rgb != RGBColor(int(color[0]), int(color[1]), int(color[2])):
            if re.search(reValue, parRuns.text) is not None:
                parRuns.font.color.rgb = RGBColor(int(color[0]), int(color[1]), int(color[2]))


# Исправление межстрочного интервала
def paragraphLineSpacingTest(parFormat, trueStyle, style, paragraph):
    if parFormat.line_spacing is None:
        if style[paragraph.style.name].paragraph_format.line_spacing is None:
            if paragraph.style.base_style is not None:
                if style[paragraph.style.base_style.name].paragraph_format.line_spacing != float(
                        replaceSymbol(trueStyle[3])):
                    style[paragraph.style.base_style.name].paragraph_format.line_spacing = float(replaceSymbol(trueStyle[3]))
        else:
            if style[paragraph.style.name].paragraph_format.line_spacing != float(replaceSymbol(trueStyle[3])):
                style[paragraph.style.name].paragraph_format.line_spacing = float(replaceSymbol(trueStyle[3]))
    else:
        if parFormat.line_spacing != float(replaceSymbol(trueStyle[3])):
            parFormat.line_spacing = float(replaceSymbol(trueStyle[3]))


# Исправление выравнивания текста
def paragraphAlignmentTest(parFormat, trueStyle, style, paragraph):
    trueValueAlignment = None
    match replaceSymbol(trueStyle[4]):
        case "Ширине":
            trueValueAlignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY
        case "Центру":
            trueValueAlignment = WD_PARAGRAPH_ALIGNMENT.CENTER
        case "Лево":
            trueValueAlignment = WD_PARAGRAPH_ALIGNMENT.LEFT
        case "Право":
            trueValueAlignment = WD_PARAGRAPH_ALIGNMENT.RIGHT
    if parFormat.alignment is None:
        if style[paragraph.style.name].paragraph_format.alignment is None:
            if paragraph.style.base_style is not None:
                if style[paragraph.style.base_style.name].paragraph_format.alignment != trueValueAlignment:
                    style[paragraph.style.base_style.name].paragraph_format.alignment = trueValueAlignment
        else:
            if style[paragraph.style.name].paragraph_format.alignment != trueValueAlignment:
                style[paragraph.style.name].paragraph_format.alignment = trueValueAlignment
    else:
        if parFormat.alignment != trueValueAlignment:
            if len(paragraph.text.replace(" ", "")) != 0:
                parFormat.alignment = trueValueAlignment


# Проверка межстрочного интервала
def paragraphFirstLineIndentTest(parFormat, trueStyle, style, paragraph):
    if parFormat.first_line_indent is None:
        if style[paragraph.style.name].paragraph_format.first_line_indent is None:
            if paragraph.style.base_style is not None:
                if round(style[paragraph.style.base_style.name].paragraph_format.first_line_indent.cm, 2) \
                        != float(replaceSymbol(trueStyle[5])):
                    style[paragraph.style.base_style.name].paragraph_format.first_line_indent = Cm(float(replaceSymbol(trueStyle[5])))
        else:
            if round(style[paragraph.style.name].paragraph_format.first_line_indent.cm, 2) != float(
                    replaceSymbol(trueStyle[5])):
                style[paragraph.style.name].paragraph_format.first_line_indent = Cm(float(replaceSymbol(trueStyle[5])))
    else:
        if round(parFormat.first_line_indent.cm, 2) != float(replaceSymbol(trueStyle[5])):
            parFormat.first_line_indent = Cm(float(replaceSymbol(trueStyle[5])))


# Проверка отступа между этим абзацем и последующим
def paragraphSpaceAfterTest(parFormat, trueStyle, style, paragraph):
    if parFormat.space_after is None:
        if style[paragraph.style.name].paragraph_format.space_after is None:
            if paragraph.style.base_style is not None:
                if style[paragraph.style.base_style.name].paragraph_format.space_after is not None:
                    if round(style[paragraph.style.base_style.name].paragraph_format.space_after.cm, 2) != float(
                            replaceSymbol(trueStyle[7])):
                        style[paragraph.style.base_style.name].paragraph_format.space_after = Cm(float(replaceSymbol(trueStyle[7])))
        else:
            if round(style[paragraph.style.name].paragraph_format.space_after.cm, 2) != float(
                    replaceSymbol(trueStyle[7])):
                style[paragraph.style.name].paragraph_format.space_after = Cm(float(replaceSymbol(trueStyle[7])))
    else:
        if round(parFormat.space_after.cm, 2) != float(replaceSymbol(trueStyle[7])):
            parFormat.space_after = Cm(float(replaceSymbol(trueStyle[7])))


# Проверка отступа между этим абзацем и предыдущим
def paragraphSpaceBeforeTest(parFormat, trueStyle, style, paragraph):
    if parFormat.space_before is None:
        if style[paragraph.style.name].paragraph_format.space_before is None:
            if paragraph.style.base_style is not None:
                if style[paragraph.style.base_style.name].paragraph_format.space_before is not None:
                    if round(style[paragraph.style.base_style.name].paragraph_format.space_before.cm, 2) != float(
                            replaceSymbol(trueStyle[8])):
                        style[paragraph.style.base_style.name].paragraph_format.space_before = Cm(float(replaceSymbol(trueStyle[8])))
        else:
            if round(style[paragraph.style.name].paragraph_format.space_before.cm, 2) != float(
                    replaceSymbol(trueStyle[8])):
                style[paragraph.style.name].paragraph_format.space_before = Cm(float(replaceSymbol(trueStyle[8])))
    else:
        if round(parFormat.space_before.cm, 2) != float(replaceSymbol(trueStyle[8])):
            parFormat.space_before = Cm(float(replaceSymbol(trueStyle[8])))


# Проверка разрыва абзаца
def paragraphKeepWithNextTest(parFormat, trueStyle, style, paragraph):
    if parFormat.keep_with_next is None:
        if style[paragraph.style.name].paragraph_format.keep_with_next is None:
            if paragraph.style.base_style is not None:
                if str(style[paragraph.style.base_style.name].paragraph_format.keep_with_next) != replaceSymbol(trueStyle[9]):
                    style[paragraph.style.base_style.name].paragraph_format.keep_with_next = strtobool(replaceSymbol(trueStyle[9]))
        else:
            if str(style[paragraph.style.name].paragraph_format.keep_with_next) != replaceSymbol(trueStyle[9]):
                 style[paragraph.style.name].paragraph_format.keep_with_next = strtobool(replaceSymbol(trueStyle[9]))
    else:
        if str(parFormat.keep_with_next) != replaceSymbol(trueStyle[9]):
            parFormat.keep_with_next = strtobool(replaceSymbol(trueStyle[9]))


# Распределение проверок, согласно имени стиля
def paragraphDesignTest(docInput, idPattern, checkTestArr):
    stylesArr = findStyleCollectionFromPattern(idPattern)
    trueFileStyleArr = [[0] * 12 for i in range(len(stylesArr) - 1)]
    count = 0
    for nameStyle in stylesArr:
        nameStyle = replaceSymbol(nameStyle)
        if nameStyle == "Common":
            if 12 in checkTestArr:
                pageMarginTest(docInput, nameStyle, idPattern)
        else:
            trueFileStyleArr[count][0] = nameStyle
            styleArr = findFileStyleCollection(nameStyle, idPattern)
            number = 1
            for style in styleArr:
                trueFileStyleArr[count][number] = style
                number += 1
            count += 1
    comparisonParagraphStyle(docInput, trueFileStyleArr, checkTestArr)


# Вызов соответствующих проверок
def comparisonParagraphStyle(docInput, trueFileStyleArr, checkTestArr):
    style = docInput.styles
    for paragraph in docInput.paragraphs:
        parFormat = paragraph.paragraph_format
        for trueStyle in trueFileStyleArr:
            if trueStyle[0] == paragraph.style.name:
                parRunsArr = paragraph.runs
                for parRuns in parRunsArr:
                    if 1 in checkTestArr:
                        fontNameTest(parRuns, trueStyle, style, paragraph)
                    if 2 in checkTestArr:
                        fontSizeTest(parRuns, trueStyle, style, paragraph)
                    if 6 in checkTestArr:
                        fontBoldItalicUnderlineTest(parRuns, trueStyle, style, paragraph)
                    if 10 in checkTestArr:
                        fontAllCapsTest(parRuns, trueStyle, style, paragraph)
                    if 11 in checkTestArr:
                        fontColorTest(parRuns, trueStyle, style, paragraph)

                if 3 in checkTestArr:
                    paragraphLineSpacingTest(parFormat, trueStyle, style, paragraph)
                if 4 in checkTestArr:
                    paragraphAlignmentTest(parFormat, trueStyle, style, paragraph)
                if 5 in checkTestArr:
                    paragraphFirstLineIndentTest(parFormat, trueStyle, style, paragraph)
                if 7 in checkTestArr:
                    paragraphSpaceAfterTest(parFormat, trueStyle, style, paragraph)
                if 8 in checkTestArr:
                    paragraphSpaceBeforeTest(parFormat, trueStyle, style, paragraph)
                if 9 in checkTestArr:
                    paragraphKeepWithNextTest(parFormat, trueStyle, style, paragraph)


# Вызов исправления файла
def startCheckAndEditFileForFormatting(fileName, checkTestArr, idPattern):
    fileExtensionOutput = ".docx"
    fileNameDocx = fileName[:fileName.rfind(".")] + fileExtensionOutput
    isConvertFile = False
    if not os.path.exists(fileNameDocx):
        isConvertFile = True
        convertFileToDOCX(fileName)

    file = open(fileNameDocx, 'rb')
    docInput = docx.Document(file)
    paragraphDesignTest(docInput, idPattern, checkTestArr)
    docInput.save('OutputFiles/Результаты проверки файла/Исправленный_' + fileName[fileName.rfind("/") + 1:fileName.rfind(".")] + ".docx")
    file.close()

    if isConvertFile:
        deleteDOCXFile(fileName)
