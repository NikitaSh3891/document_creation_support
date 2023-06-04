import docx
from docx.shared import Pt
import os
import re

from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from thefuzz import fuzz
from docx.shared import RGBColor

from Converter import convertFileToDOCX, deleteDOCXFile
from WorkWithDB import findStyleCollectionFromPattern, findFileStyleCollection
from WorkWithFileCreate.EditFileStyle import replaceSymbol
from WorkWithFileCreate.CreateСommentInFile import findDisplayName, addComment

"""
Класс предназначен для проверки файлов на форматирование
"""


reValue = "[a-z]|[A-Z]|[а-я]|[А-Я]|[ё|Ё]"
dictionaryOfQuotationMarks = ["'", "\"", "«", "»"]


# Фильтрация элементов по их типу
def filterElement(document):
    res = []
    for element in document.elements:
        if 'paragraph' in str(element):
            res.append(element)
        elif 'table' in str(element):
            for row in element.rows:
                for cell in row.cells:
                    res.append(cell.paragraphs)
    return res


# Определение номера страницы переданного абзаца (работает с небольшой погрешностью)
def findNumberPage(text, paragraphs):
    # список для сбора номеров страниц, где встретилась строка
    pages = []
    # счетчик для номеров страниц
    numberPage = 1
    for paragraph in paragraphs:
        if type(paragraph) == list:
            # проверка на наличие разрывов страниц внутри таблиц
            text_in_table = [p.text for p in paragraph]
            text_in_table = ''.join(text_in_table)
            for p in paragraph:
                for run in p.runs:
                    # проверка на мягкий разрыва страницы
                    if 'lastRenderedPageBreak' in run._element.xml:
                        numberPage += 1
                    # проверка на жесткий разрыв страницы
                    elif 'w:br' in run._element.xml and 'type="page"' in run._element.xml:
                        numberPage += 1
            # к этому моменту известен актуальный номер страницы
            if len(text_in_table) >= len(text) - 10:
                res = fuzz.partial_ratio(text.lower(), text_in_table.lower())
                if res >= 97:
                    # если строчка найдена, добавляем в список номер страницы
                    pages.append(numberPage)
        else:
            # проверка разрывов в абзацах вне таблиц
            for run in paragraph.runs:
                if 'lastRenderedPageBreak' in run._element.xml:
                    numberPage += 1
                elif 'w:br' in run._element.xml and 'type="page"' in run._element.xml:
                    numberPage += 1
            if len(paragraph.text) >= len(text):
                res = fuzz.partial_ratio(text.lower(), paragraph.text.lower())
                if res >= 97:
                    pages.append(numberPage)
    return ', '.join(map(str, pages))


# Число букв, выводимое в файл для определения места ошибки
def editNumCharacter(numCharacter):
    global num
    num = numCharacter


# Исправление отступов документа
def pageMarginTest(docInput, nameStyle, docOutput, idPattern):
    styleArr = findFileStyleCollection(nameStyle, idPattern)
    sections = docInput.sections
    countSections = 0
    for section in sections:
        if round(section.left_margin.cm, 2) != float(replaceSymbol(styleArr[0])):
            docOutput.write(f"[IdErr-12] Левые отступы {countSections} раздела документа не соблюдены!\n")
        if round(section.right_margin.cm, 2) != float(replaceSymbol(styleArr[1])):
            docOutput.write(f"[IdErr-12] Правые отступы {countSections} раздела документа не соблюдены!\n")
        if round(section.top_margin.cm, 2) != float(replaceSymbol(styleArr[2])):
            docOutput.write(f"[IdErr-12] Верхние отступы {countSections} раздела документа не соблюдены!\n")
        if round(section.bottom_margin.cm, 2) != float(replaceSymbol(styleArr[3])):
            docOutput.write(f"[IdErr-12] Нижние отступы {countSections} раздела документа не соблюдены!\n")
        countSections += 1


# Проверка шрифта
def fontNameTest(parRuns, trueStyle, style,  paragraph, paragraphs, docOutput):
    if parRuns.font.name is None:
        if style[paragraph.style.name].font.name != replaceSymbol(trueStyle[1]):
            if paragraph.style.base_style is not None:
                if style[paragraph.style.base_style.name].font.name is not None:
                    if style[paragraph.style.base_style.name].font.name != replaceSymbol(trueStyle[1]):
                        text = parRuns.text[:num].translate({ord(i): None for i in '\n'})
                        if len(text.replace(" ", "")) > 0:
                            addMsgAboutErr(paragraph, docOutput, paragraphs, f"шрифте в части: \"{text}\"", 1)
    else:
        if parRuns.font.name != replaceSymbol(trueStyle[1]):
            text = parRuns.text[:num].translate({ord(i): None for i in '\n'})
            if len(text.replace(" ", "")) > 0:
                if parRuns.font.name == "Arial":
                    addMsgAboutErr(paragraph, docOutput, paragraphs, f"шрифте (не точно) в части: \"{text}\"", 1)
                else:
                    addMsgAboutErr(paragraph, docOutput, paragraphs, f"шрифте в части: \"{text}\"", 1)


# Проверка размера шрифта
def fontSizeTest(parRuns, trueStyle, style, paragraph, paragraphs, docOutput):
    if parRuns.font.size is None:
        if style[paragraph.style.name].font.size != Pt(int(replaceSymbol(trueStyle[2]))):
            if paragraph.style.base_style is not None:
                if style[paragraph.style.base_style.name].font.size is not None:
                    if style[paragraph.style.base_style.name].font.size != Pt(int(replaceSymbol(trueStyle[2]))):
                        text = parRuns.text[:num].translate({ord(i): None for i in '\n'})
                        if len(text.replace(" ", "")) > 0:
                            addMsgAboutErr(paragraph, docOutput, paragraphs, f"размере шрифта в части: \"{text}\"", 2)
    else:
        if parRuns.font.size != Pt(int(replaceSymbol(trueStyle[2]))):
            text = parRuns.text[:num].translate({ord(i): None for i in '\n'})
            if len(text.replace(" ", "")) > 0:
                addMsgAboutErr(paragraph, docOutput, paragraphs, f"размере шрифта в части: \"{text}\"", 2)


# Проверка на полужирный, подчеркнутый и курсивный текст
def fontBoldItalicUnderlineTest(parRuns, trueStyle, style, paragraph, paragraphs, docOutput):
    trueValueBold = None
    trueValueItalic = None
    trueValueUnderline = None
    match replaceSymbol(trueStyle[6]):
        case "Жирный":
            trueValueBold = True
            trueValueItalic = False
            trueValueUnderline = False
        case "Курсив":
            trueValueBold = False
            trueValueItalic = True
            trueValueUnderline = False
        case "Подчеркивание":
            trueValueBold = False
            trueValueItalic = False
            trueValueUnderline = True
        case "ЖирныйИКурсив":
            trueValueBold = True
            trueValueItalic = True
            trueValueUnderline = False
        case "ЖирныйИПодчеркивание":
            trueValueBold = True
            trueValueItalic = False
            trueValueUnderline = True
        case "КурсивИПодчеркивание":
            trueValueBold = False
            trueValueItalic = True
            trueValueUnderline = True
        case "ЖирныйИКурсивИПодчеркивание":
            trueValueBold = True
            trueValueItalic = True
            trueValueUnderline = True
        case "Обычный":
            trueValueBold = False
            trueValueItalic = False
            trueValueUnderline = False
    if parRuns.font.bold is None:
        if style[paragraph.style.name].font.bold and not trueValueBold:
            if len(parRuns.text.replace(" ", "")) > 0:
                addMsgAboutErr(paragraph, docOutput, paragraphs, "полужирном начертании текста", 6)
    else:
        if parRuns.font.bold and not trueValueBold:
            text = parRuns.text[:num].translate({ord(i): None for i in '\n'})
            if len(text.replace(" ", "")) > 0:
                addMsgAboutErr(paragraph, docOutput, paragraphs, f"полужирном начертании текста в части: \"{text}\"", 6)
    if parRuns.font.italic is None:
        if style[paragraph.style.name].font.italic and not trueValueItalic:
            if len(parRuns.text.replace(" ", "")) > 0:
                addMsgAboutErr(paragraph, docOutput, paragraphs, "курсивном начертании текста", 6)
    else:
        if parRuns.font.italic and not trueValueItalic:
            text = parRuns.text[:num].translate({ord(i): None for i in '\n'})
            if len(text.replace(" ", "")) > 0:
                addMsgAboutErr(paragraph, docOutput, paragraphs, f"курсивном начертании текста в части: \"{text}\"", 6)
    if parRuns.font.underline is None:
        if style[paragraph.style.name].font.underline and not trueValueUnderline:
            if len(parRuns.text.replace(" ", "")) > 0:
                addMsgAboutErr(paragraph, docOutput, paragraphs, "подчеркнутом начертании текста", 6)
    else:
        if parRuns.font.underline and not trueValueUnderline:
            text = parRuns.text[:num].translate({ord(i): None for i in '\n'})
            if len(text.replace(" ", "")) > 0:
                addMsgAboutErr(paragraph, docOutput, paragraphs, f"подчеркнутом начертании текста в части: \"{text}\"", 6)


# Проверка на заглавные буквы
def fontAllCapsTest(parRuns, trueStyle, style, paragraph, paragraphs, docOutput):
    if parRuns.font.all_caps is None:
        if replaceSymbol(trueStyle[10]) == "False":
            if parRuns.text == parRuns.text.upper():
                if re.search(reValue, parRuns.text) is not None:
                    if len(re.sub('[0-9 ]+', '', parRuns.text)) != 1:
                        numberStartStr = paragraph.text.find(parRuns.text)
                        if paragraph.text[numberStartStr - 1] != " " and \
                                re.search(reValue, paragraph.text[numberStartStr - 1]) is None:
                            if paragraph.text[numberStartStr - 1] not in dictionaryOfQuotationMarks:
                                text = parRuns.text[:num].translate({ord(i): None for i in '\n'})
                                addMsgAboutErr(paragraph, docOutput, paragraphs, f"заглавных буквах текста в части: \"{text}\"", 10)
                        else:
                            isTrue = False
                            numSpace = numberStartStr - 1
                            while numSpace > 0:
                                if paragraph.text[numSpace] in dictionaryOfQuotationMarks:
                                    isTrue = True
                                    break
                                numSpace -= 1
                            if not isTrue:
                                text = parRuns.text[:num].translate({ord(i): None for i in '\n'})
                                addMsgAboutErr(paragraph, docOutput, paragraphs, f"заглавных буквах текста в части: \"{text}\"", 10)
    else:
        if str(parRuns.font.all_caps) != replaceSymbol(trueStyle[10]):
            if re.search(reValue, parRuns.text) is not None:
                text = parRuns.text[:num].translate({ord(i): None for i in '\n'})
                addMsgAboutErr(paragraph, docOutput, paragraphs, f"заглавных буквах текста в части: \"{text}\"", 10)


# Проверка на цвет текста
def fontColorTest(parRuns, trueStyle, style, paragraph, paragraphs, docOutput):
    if parRuns.font.color.rgb is not None:
        color = replaceSymbol(trueStyle[11]).split(" ")
        if parRuns.font.color.rgb != RGBColor(int(color[0]), int(color[1]), int(color[2])):
            if re.search(reValue, parRuns.text) is not None:
                text = parRuns.text[:num].translate({ord(i): None for i in '\n'})
                addMsgAboutErr(paragraph, docOutput, paragraphs, f"цвете текста в части: \"{text}\"", 11)


# Проверка межстрочного интервала
def paragraphLineSpacingTest(parFormat, trueStyle, style, paragraph, paragraphs, docOutput):
    # Значение, указывающее расстояние между базовыми линиями в последовательных строках абзаца.
    # Значение None указывает, что межстрочный интервал наследуется из иерархии стилей.
    # Значение None у определенного стиля указывает, что оно наследуется от стиля Normal
    if parFormat.line_spacing is None:
        if style[paragraph.style.name].paragraph_format.line_spacing is None:
            if paragraph.style.base_style is not None:
                if style[paragraph.style.base_style.name].paragraph_format.line_spacing != float(
                        replaceSymbol(trueStyle[3])):
                    addMsgAboutErr(paragraph, docOutput, paragraphs, "межстрочном интервале", 3)
        else:
            if style[paragraph.style.name].paragraph_format.line_spacing != float(replaceSymbol(trueStyle[3])):
                addMsgAboutErr(paragraph, docOutput, paragraphs, "межстрочном интервале", 3)
    else:
        if parFormat.line_spacing != float(replaceSymbol(trueStyle[3])):
            addMsgAboutErr(paragraph, docOutput, paragraphs, "межстрочном интервале", 3)


# Проверка выравнивания текста
def paragraphAlignmentTest(parFormat, trueStyle, style, paragraph, paragraphs, docOutput):
    trueValueAlignment = None
    match replaceSymbol(trueStyle[4]):
        case "Ширине":
            trueValueAlignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY
        case "Центру":
            trueValueAlignment = WD_PARAGRAPH_ALIGNMENT.CENTER
        case "Лево":
            trueValueAlignment = WD_PARAGRAPH_ALIGNMENT.LEFT
        case "Право":
            trueValueAlignment = WD_PARAGRAPH_ALIGNMENT.RIGHT
    if parFormat.alignment is None:
        if style[paragraph.style.name].paragraph_format.alignment is None:
            if paragraph.style.base_style is not None:
                if style[paragraph.style.base_style.name].paragraph_format.alignment != trueValueAlignment:
                    addMsgAboutErr(paragraph, docOutput, paragraphs, "выравнивании текста", 4)
        else:
            if style[paragraph.style.name].paragraph_format.alignment != trueValueAlignment:
                addMsgAboutErr(paragraph, docOutput, paragraphs, "выравнивании текста", 4)
    else:
        if parFormat.alignment != trueValueAlignment:
            if len(paragraph.text.replace(" ", "")) > 0:
                addMsgAboutErr(paragraph, docOutput, paragraphs, "выравнивании текста", 4)


# Проверка отступа первой строки
def paragraphFirstLineIndentTest(parFormat, trueStyle, style, paragraph, paragraphs, docOutput):
    if parFormat.first_line_indent is None:
        if style[paragraph.style.name].paragraph_format.first_line_indent is None:
            if paragraph.style.base_style is not None:
                if round(style[paragraph.style.base_style.name].paragraph_format.first_line_indent.cm, 2) \
                        != float(replaceSymbol(trueStyle[5])):
                    addMsgAboutErr(paragraph, docOutput, paragraphs, "отступе первой строки", 5)
        else:
            if round(style[paragraph.style.name].paragraph_format.first_line_indent.cm, 2) != float(
                    replaceSymbol(trueStyle[5])):
                addMsgAboutErr(paragraph, docOutput, paragraphs, "отступе первой строки", 5)
    else:
        if round(parFormat.first_line_indent.cm, 2) != float(replaceSymbol(trueStyle[5])):
            addMsgAboutErr(paragraph, docOutput, paragraphs, "отступе первой строки", 5)


# Проверка отступа между этим абзацем и последующим
def paragraphSpaceAfterTest(parFormat, trueStyle, style, paragraph, paragraphs, docOutput):
    if parFormat.space_after is None:
        if style[paragraph.style.name].paragraph_format.space_after is None:
            if paragraph.style.base_style is not None:
                if style[paragraph.style.base_style.name].paragraph_format.space_after is not None:
                    if round(style[paragraph.style.base_style.name].paragraph_format.space_after.cm, 2) != float(
                            replaceSymbol(trueStyle[7])):
                        addMsgAboutErr(paragraph, docOutput, paragraphs, "верхнем отступе", 7)
        else:
            if round(style[paragraph.style.name].paragraph_format.space_after.cm, 2) != float(
                    replaceSymbol(trueStyle[7])):
                addMsgAboutErr(paragraph, docOutput, paragraphs, "верхнем отступе", 7)
    else:
        if round(parFormat.space_after.cm, 2) != float(replaceSymbol(trueStyle[7])):
            addMsgAboutErr(paragraph, docOutput, paragraphs, "верхнем отступе", 7)


# Проверка отступа между этим абзацем и предыдущим
def paragraphSpaceBeforeTest(parFormat, trueStyle, style, paragraph, paragraphs, docOutput):
    if parFormat.space_before is None:
        if style[paragraph.style.name].paragraph_format.space_before is None:
            if paragraph.style.base_style is not None:
                if style[paragraph.style.base_style.name].paragraph_format.space_before is not None:
                    if round(style[paragraph.style.base_style.name].paragraph_format.space_before.cm, 2) != float(
                            replaceSymbol(trueStyle[8])):
                        addMsgAboutErr(paragraph, docOutput, paragraphs, "нижнем отступе", 8)
        else:
            if round(style[paragraph.style.name].paragraph_format.space_before.cm, 2) != float(
                    replaceSymbol(trueStyle[8])):
                addMsgAboutErr(paragraph, docOutput, paragraphs, "нижнем отступе", 8)
    else:
        if round(parFormat.space_before.cm, 2) != float(replaceSymbol(trueStyle[8])):
            addMsgAboutErr(paragraph, docOutput, paragraphs, "нижнем отступе", 8)


# Проверка разрыва абзаца
def paragraphKeepWithNextTest(parFormat, trueStyle, style, paragraph, paragraphs, docOutput):
    if parFormat.keep_with_next is None:
        if style[paragraph.style.name].paragraph_format.keep_with_next is None:
            if paragraph.style.base_style is not None:
                if str(style[paragraph.style.base_style.name].paragraph_format.keep_with_next) != replaceSymbol(trueStyle[9]):
                    addMsgAboutErr(paragraph, docOutput, paragraphs, "разрыве абзаца", 9)
        else:
            if str(style[paragraph.style.name].paragraph_format.keep_with_next) != replaceSymbol(trueStyle[9]):
                addMsgAboutErr(paragraph, docOutput, paragraphs, "разрыве абзаца", 9)
    else:
        if str(parFormat.keep_with_next) != replaceSymbol(trueStyle[9]):
            addMsgAboutErr(paragraph, docOutput, paragraphs, "разрыве абзаца", 9)


# Распределение проверок, согласно имени стиля
def paragraphDesignTest(docInput, idPattern, docOutput, checkTestArr):
    stylesArr = findStyleCollectionFromPattern(idPattern)
    trueFileStyleArr = [[0] * 12 for i in range(len(stylesArr) - 1)]
    count = 0
    for nameStyle in stylesArr:
        nameStyle = replaceSymbol(nameStyle)
        if nameStyle == "Common":
            if 12 in checkTestArr:
                pageMarginTest(docInput, nameStyle, docOutput, idPattern)
        else:
            trueFileStyleArr[count][0] = nameStyle
            styleArr = findFileStyleCollection(nameStyle, idPattern)
            number = 1
            for style in styleArr:
                trueFileStyleArr[count][number] = style
                number += 1
            count += 1
    comparisonParagraphStyle(docInput, trueFileStyleArr, docOutput, checkTestArr)


# Вызов соответствующих проверок
def comparisonParagraphStyle(docInput, trueFileStyleArr, docOutput, checkTestArr):
    style = docInput.styles
    paragraphs = filterElement(docInput)
    global nameUser
    nameUser = findDisplayName()
    for paragraph in docInput.paragraphs:
        parFormat = paragraph.paragraph_format
        for trueStyle in trueFileStyleArr:
            if trueStyle[0] == paragraph.style.name:
                parRunsArr = paragraph.runs
                for parRuns in parRunsArr:
                    if 1 in checkTestArr:
                        fontNameTest(parRuns, trueStyle, style, paragraph, paragraphs, docOutput)
                    if 2 in checkTestArr:
                        fontSizeTest(parRuns, trueStyle, style, paragraph, paragraphs, docOutput)
                    if 6 in checkTestArr:
                        fontBoldItalicUnderlineTest(parRuns, trueStyle, style, paragraph, paragraphs, docOutput)
                    if 10 in checkTestArr:
                        fontAllCapsTest(parRuns, trueStyle, style, paragraph, paragraphs, docOutput)
                    if 11 in checkTestArr:
                        fontColorTest(parRuns, trueStyle, style, paragraph, paragraphs, docOutput)

                if 3 in checkTestArr:
                    paragraphLineSpacingTest(parFormat, trueStyle, style, paragraph, paragraphs, docOutput)
                if 4 in checkTestArr:
                    paragraphAlignmentTest(parFormat, trueStyle, style, paragraph, paragraphs, docOutput)
                if 5 in checkTestArr:
                    paragraphFirstLineIndentTest(parFormat, trueStyle, style, paragraph, paragraphs, docOutput)
                if 7 in checkTestArr:
                    paragraphSpaceAfterTest(parFormat, trueStyle, style, paragraph, paragraphs, docOutput)
                if 8 in checkTestArr:
                    paragraphSpaceBeforeTest(parFormat, trueStyle, style, paragraph, paragraphs, docOutput)
                if 9 in checkTestArr:
                    paragraphKeepWithNextTest(parFormat, trueStyle, style, paragraph, paragraphs, docOutput)


# Запись ошибок в файл .txt
def addMsgAboutErr(paragraph, docOutput, paragraphs, msgErr, idErr):
    num = 10
    if len(paragraph.text) < num:
        num = len(paragraph.text)
    # Удаление из строки всех знаков перехода на новую строку (\n)
    text = paragraph.text[:num].translate({ord(i): None for i in '\n'})
    if len(text.replace(" ", "")) > 0:
        pageNumber = findNumberPage(text, paragraphs)
        pageNumArr = pageNumber.replace(" ", "").split(",")
        temp = []
        for i in pageNumArr:
            if i not in temp:
                temp.append(i)
        pageNumber = ', '.join(temp)

        try:
            addComment("Ошибка в " + msgErr, paragraph.text, paragraphs, nameUser)
            docOutput.write(f"[IdErr-{idErr}] На стр. {pageNumber} содержится ошибка в " +
                            msgErr + f" в абзаце: \"{text}\"\n")
        except UnicodeEncodeError:
            docOutput.write(f"[IdErr-{idErr}] На стр. {pageNumber} содержится ошибка в " +
                            replaceStressOnText(msgErr) + f" в абзаце: \"{replaceStressOnText(text)}\"\n")


# Из-за ошибки кодировки необходимо заменять буквы с ударением на обычные буквы
def replaceStressOnText(text):
    return text.replace("А́", "").replace("а́", "").replace("а́", "").replace("Е́", "").replace("е́", "").replace("И́", "").replace(
        "и́", "").replace("О́", "").replace("о́", "").replace("У́", "").replace("у́", "").replace("ы́", "").replace("Э́", "").replace(
        "э́", "").replace("Ю́", "").replace("ю́", "").replace("Я́", "").replace("я́", "")


# Вызов проверки файла
def startCheckForFileFormatting(fileName, checkTestArr, idPattern):
    fileExtensionOutput = ".docx"
    fileNameDocx = fileName[:fileName.rfind(".")] + fileExtensionOutput
    isConvertFile = False
    if not os.path.exists(fileNameDocx):
        isConvertFile = True
        convertFileToDOCX(fileName)

    file = open(fileNameDocx, 'rb')
    import zipfile
    try:
        docInput = docx.Document(file)
        docOutput = open('OutputFiles/Результаты проверки файла/' + fileName[fileName.rfind("/") + 1:fileName.rfind(".")] + ".txt", "w")
        paragraphDesignTest(docInput, idPattern, docOutput, checkTestArr)
        docInput.save('OutputFiles/Результаты проверки файла/Комментарии_' + fileName[fileName.rfind("/") + 1:fileName.rfind(".")] + ".docx")
        docOutput.close()
        file.close()

        if isConvertFile:
            deleteDOCXFile(fileName)
    except zipfile.BadZipFile:
        from Windows.QtWindows.errWindow import editText
        editText(f"Ошибка! Файл {fileNameDocx} пустой")
